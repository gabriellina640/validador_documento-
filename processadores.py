import docx
import openpyxl
import fitz 

def processar_word(caminho_entrada, caminho_saida, dados_auditoria, hash_original):
    try:
        doc = docx.Document(caminho_entrada)
        doc.add_paragraph("\n" + "="*50)
        p_titulo = doc.add_paragraph("REGISTRO DE AUTENTICIDADE - CORREGEDORIA")
        p_titulo.runs[0].bold = True
        doc.add_paragraph(f"🕒 Data/Hora: {dados_auditoria['data_hora']}")
        doc.add_paragraph(f"👤 Operador: {dados_auditoria['usuario_sistema']}")
        doc.add_paragraph(f"🔐 Autenticidade (SHA-256): {hash_original}") 
        doc.add_paragraph("="*50)
        doc.save(caminho_saida)
        return True
    except Exception as e: 
        print(f"Erro no Word: {e}")
        return False

def processar_excel(caminho_entrada, caminho_saida, dados_auditoria, hash_original):
    try:
        wb = openpyxl.load_workbook(caminho_entrada)
        nome_aba = "Autenticidade_Corregedoria"
        if nome_aba not in wb.sheetnames: 
            ws = wb.create_sheet(nome_aba)
        else: 
            ws = wb[nome_aba]
        
        ws.append(["REGISTRO DE AUTENTICIDADE", ""])
        ws.append(["Data/Hora", dados_auditoria['data_hora']])
        ws.append(["Operador", dados_auditoria['usuario_sistema']])
      
        ws.append(["Hash SHA-256", hash_original]) 
        ws.append([""]) 
        
        
        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 80
        
        wb.save(caminho_saida)
        return True
    except Exception as e: 
        print(f"Erro no Excel: {e}")
        return False

def processar_pdf(caminho_entrada, caminho_saida, dados_auditoria, hash_original):
    try:
        doc = fitz.open(caminho_entrada)
        
        
        linha1 = f"AUTENTICADO - CORREGEDORIA | {dados_auditoria['data_hora']} | Operador: {dados_auditoria['usuario_sistema']}"
        linha2 = f"Chave de Validação (SHA-256): {hash_original}"
        
        for pagina in doc:
            retangulo = pagina.rect
            ponto_linha1 = fitz.Point(30, retangulo.height - 30)
            ponto_linha2 = fitz.Point(30, retangulo.height - 20)
            
            pagina.insert_text(ponto_linha1, linha1, fontsize=7, color=(0.8, 0, 0))
            pagina.insert_text(ponto_linha2, linha2, fontsize=7, color=(0.8, 0, 0))
            
        doc.save(caminho_saida)
        return True
    except Exception as e: 
        print(f"Erro no PDF: {e}")
        return False