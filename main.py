import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import datetime
import docx

import seguranca
import banco_dados
import processadores

ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue") 

class AppCorregedoria(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        banco_dados.inicializar_banco()

        self.title("Sistema de Validação Institucional - MPAC")
        self.geometry("700x550")
        
        try:
            
            self.iconbitmap("icone.ico")
        except:
            pass 

        self.lbl_titulo = ctk.CTkLabel(self, text="🏛️ Corregedoria - Auditoria de Documentos", font=ctk.CTkFont(size=20, weight="bold"))
        self.lbl_titulo.pack(pady=(20, 10))

        self.abas = ctk.CTkTabview(self, width=650, height=450)
        self.abas.pack(padx=20, pady=10, fill="both", expand=True)

        self.aba_processar = self.abas.add(" ⚙️ Processar e Carimbar ")
        self.aba_auditar = self.abas.add(" 🔍 Verificar Autenticidade ")

        self.montar_aba_processar()
        self.montar_aba_auditar()
        
        self.ultimo_resultado_auditoria = None
        self.ultimo_arquivo_suspeito = ""

    def montar_aba_processar(self):
        lbl_desc = ctk.CTkLabel(self.aba_processar, text="Selecione os rascunhos originais para aplicar o selo oficial de validação.", font=ctk.CTkFont(size=14))
        lbl_desc.pack(pady=20)

        btn_selecionar = ctk.CTkButton(self.aba_processar, text="Selecionar Arquivos", command=self.processar_arquivos, height=40, font=ctk.CTkFont(size=14, weight="bold"))
        btn_selecionar.pack(pady=10)

        self.caixa_logs = ctk.CTkTextbox(self.aba_processar, width=600, height=200)
        self.caixa_logs.pack(pady=20)
        self.caixa_logs.insert("0.0", "Aguardando seleção de arquivos...\n")
        self.caixa_logs.configure(state="disabled")

    def log_mensagem(self, mensagem):
        self.caixa_logs.configure(state="normal")
        self.caixa_logs.insert("end", mensagem + "\n")
        self.caixa_logs.see("end")
        self.caixa_logs.configure(state="disabled")
        self.update()

    def processar_arquivos(self):
        arquivos = filedialog.askopenfilenames(filetypes=[("Suportados", ("*.pdf", "*.docx", "*.xlsx")), ("Todos", "*.*")])
        if not arquivos: return

        hoje = datetime.datetime.now()
        pasta_destino = os.path.join(os.getcwd(), "Arquivos_Processados_Corregedoria", hoje.strftime("%Y"), hoje.strftime("%m"), hoje.strftime("%d"))
        os.makedirs(pasta_destino, exist_ok=True)

        self.log_mensagem(f"Iniciando processamento. Destino: {pasta_destino}")
        sucesso_count = erro_count = 0

        for cam_original in arquivos:
            nome_arq = os.path.basename(cam_original)
            nome_sem_ext, ext = os.path.splitext(nome_arq)
            cam_saida = os.path.join(pasta_destino, f"{nome_sem_ext}_PROCESSADO{ext}")
            
            dados = seguranca.coletar_dados_auditoria()
            hash_orig = seguranca.gerar_hash_arquivo(cam_original)
            
            sucesso = False
            ext_l = ext.lower()
            if ext_l == '.docx': sucesso = processadores.processar_word(cam_original, cam_saida, dados, hash_orig)
            elif ext_l == '.xlsx': sucesso = processadores.processar_excel(cam_original, cam_saida, dados, hash_orig)
            elif ext_l == '.pdf': sucesso = processadores.processar_pdf(cam_original, cam_saida, dados, hash_orig)
                
            if sucesso:
                hash_final = seguranca.gerar_hash_arquivo(cam_saida)
                banco_dados.registrar_documento(nome_arq, cam_saida, hash_final, dados['usuario_sistema'], dados['data_hora'])
                self.log_mensagem(f"✅ Sucesso: {nome_arq}")
                sucesso_count += 1
            else:
                self.log_mensagem(f"❌ Erro: {nome_arq}")
                erro_count += 1

        messagebox.showinfo("Concluído", f"Processamento Finalizado!\n\nSalvos em:\n{pasta_destino}")

    def montar_aba_auditar(self):
        lbl_desc = ctk.CTkLabel(self.aba_auditar, text="ATENÇÃO: Selecione um documento OFICIAL (já carimbado)\npara verificar se ele sofreu adulterações após a emissão.", font=ctk.CTkFont(size=14), text_color="#FF9800")
        lbl_desc.pack(pady=15)

        btn_verificar = ctk.CTkButton(self.aba_auditar, text="Selecionar Arquivo para Auditoria", command=self.verificar_arquivo, fg_color="#FF9800", hover_color="#F57C00", height=40, font=ctk.CTkFont(size=14, weight="bold"))
        btn_verificar.pack(pady=10)

        self.frame_res = ctk.CTkFrame(self.aba_auditar, width=600, height=180)
        self.frame_res.pack(pady=15, padx=20, fill="x")

        self.lbl_res_status = ctk.CTkLabel(self.frame_res, text="Aguardando arquivo...", font=ctk.CTkFont(size=16, weight="bold"))
        self.lbl_res_status.pack(pady=10)
        self.lbl_res_detalhes = ctk.CTkLabel(self.frame_res, text="", justify="left")
        self.lbl_res_detalhes.pack(pady=5, padx=10, anchor="w")

        self.btn_laudo = ctk.CTkButton(self.aba_auditar, text="📄 Gerar Laudo Oficial (Word)", command=self.gerar_laudo_oficial, fg_color="#1976D2", hover_color="#1565C0", height=40)
        self.btn_laudo.pack(pady=10)
        self.btn_laudo.configure(state="disabled")

    def verificar_arquivo(self):
        caminho = filedialog.askopenfilename(title="Selecione o arquivo suspeito")
        if not caminho: return
        
        self.ultimo_arquivo_suspeito = caminho
        hash_calc = seguranca.gerar_hash_arquivo(caminho)
        resultado = banco_dados.buscar_por_hash(hash_calc)
        
        self.btn_laudo.configure(state="normal") 

        if resultado:
            self.ultimo_resultado_auditoria = {"status": "AUTENTICO", "dados": resultado, "hash": hash_calc}
            self.lbl_res_status.configure(text="✅ DOCUMENTO AUTÊNTICO", text_color="#4CAF50")
            self.lbl_res_detalhes.configure(text=f"Nome Original: {resultado[0]}\nProcessado em: {resultado[1]}\nOperador: {resultado[2]}\nHash: {hash_calc[:30]}...")
        else:
            self.ultimo_resultado_auditoria = {"status": "FRAUDADO", "dados": None, "hash": hash_calc}
            self.lbl_res_status.configure(text="❌ DOCUMENTO ADULTERADO / NÃO OFICIAL", text_color="#F44336")
            self.lbl_res_detalhes.configure(text=f"Aviso: O conteúdo deste arquivo não confere com o banco de dados.\nIsso ocorre se o documento foi alterado ou se você inseriu\num rascunho sem o carimbo oficial da Corregedoria.\nHash Encontrado: {hash_calc[:30]}...")

    def gerar_laudo_oficial(self):
        if not self.ultimo_resultado_auditoria: return
        
        dados = seguranca.coletar_dados_auditoria()
        nome_arq_suspeito = os.path.basename(self.ultimo_arquivo_suspeito)
        status = self.ultimo_resultado_auditoria["status"]
        
        doc = docx.Document()
        doc.add_heading('LAUDO TÉCNICO DE AUDITORIA DE DOCUMENTO - MPAC', 0)
        
        doc.add_paragraph(f"Data da Emissão do Laudo: {dados['data_hora']}")
        doc.add_paragraph(f"Auditor Responsável: {dados['usuario_sistema']}")
        doc.add_paragraph(f"Arquivo Inspecionado: {nome_arq_suspeito}")
        doc.add_paragraph(f"Assinatura Criptográfica (SHA-256) do Arquivo Inspecionado:\n{self.ultimo_resultado_auditoria['hash']}\n")
        
        p_resultado = doc.add_paragraph()
        p_resultado.add_run("CONCLUSÃO DA AUDITORIA: ").bold = True
        
        if status == "AUTENTICO":
            info = self.ultimo_resultado_auditoria["dados"]
            p_resultado.add_run("DOCUMENTO AUTÊNTICO E ÍNTEGRO.").bold = True
            doc.add_paragraph("Atesta-se que a assinatura criptográfica do documento inspecionado confere exatamente com o registro oficial no banco de dados da instituição.")
            doc.add_paragraph(f"Histórico Oficial:\n- Processado originalmente em: {info[1]}\n- Pelo operador: {info[2]}")
        else:
            p_resultado.add_run("DOCUMENTO ADULTERADO / NÃO RECONHECIDO.").bold = True
            doc.add_paragraph("Atesta-se que a assinatura criptográfica do documento inspecionado NÃO confere com nenhum registro válido no banco de dados. O arquivo sofreu alterações de conteúdo após sua possível emissão original, não possui o selo institucional ou foi forjado.")

        pasta_laudos = os.path.join(os.getcwd(), "Laudos_Corregedoria")
        os.makedirs(pasta_laudos, exist_ok=True)
        caminho_laudo = os.path.join(pasta_laudos, f"Laudo_Auditoria_{status}_{datetime.datetime.now().strftime('%H%M%S')}.docx")
        
        try:
            doc.save(caminho_laudo)
            messagebox.showinfo("Laudo Gerado", f"Laudo oficial gerado com sucesso!\nSalvo na pasta:\n{caminho_laudo}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao gerar laudo: {e}")

if __name__ == "__main__":
    app = AppCorregedoria()
    app.mainloop()